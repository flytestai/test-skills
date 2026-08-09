#!/usr/bin/env python3
"""
Test Report DOCX Generator (测试报告生成器)

Generates a Word (.docx) test report from a JSON data file, matching the format
of the standard template: 测试报告-PC 2.1.7版本测试.docx

Template format:
- Paragraph-based (no tables)
- Section headers: List Paragraph, 18pt bold 宋体
- Content: List Paragraph, 14pt 宋体
- Scope items: 12pt 宋体

Usage:
    python generate_report.py --input data.json --output report.docx
    python generate_report.py --output report.docx --json-stdin
"""

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("[ERROR] python-docx is required. Install it with: pip install python-docx")
    sys.exit(1)


# ── helpers ──────────────────────────────────────────────────────────

def _east(run, font_name="宋体"):
    """Set East-Asian font on a run."""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_section_header(doc, text, size=18, bold=True):
    """Add a section header paragraph (18pt bold 宋体, List Paragraph style)."""
    p = doc.add_paragraph()
    p.style = doc.styles["List Paragraph"]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    _east(run, "宋体")
    return p


def add_content(doc, text, size=14, bold=False):
    """Add a content paragraph (14pt 宋体, List Paragraph style)."""
    p = doc.add_paragraph()
    p.style = doc.styles["List Paragraph"]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    _east(run, "宋体")
    return p


def add_blank(doc):
    """Add an empty paragraph for spacing."""
    p = doc.add_paragraph()
    p.style = doc.styles["List Paragraph"]
    return p


# ── section builders ─────────────────────────────────────────────────

def build_title(doc, meta):
    """Title: PC X.X.X版本测试测试报告 (Heading 1, centered)."""
    platform = meta.get("platform", "PC")
    version = meta.get("version", "X.X.X")
    title_text = f"{platform} {version}版本测试测试报告"

    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title_text)
    _east(run, "宋体")
    return p


def build_requirement_url(doc, meta):
    """Section: 需求地址"""
    add_section_header(doc, "需求地址")
    url = meta.get("requirement_url", "")
    if url:
        add_content(doc, url)
    else:
        add_content(doc, "（未提供）")
    add_blank(doc)


def build_test_scope(doc, data):
    """Section: 测试范围 (numbered list, 12pt)."""
    add_section_header(doc, "测试范围")
    scope_items = data.get("test_scope", [])
    if not scope_items:
        add_content(doc, "（无）", size=12)
        add_blank(doc)
        return

    # Build numbered list in a single paragraph with line breaks
    p = doc.add_paragraph()
    p.style = doc.styles["List Paragraph"]

    for i, item in enumerate(scope_items, 1):
        run = p.add_run(f"{i}.{item}")
        run.font.size = Pt(12)
        _east(run, "宋体")
        if i < len(scope_items):
            # Add line break between items
            run_br = p.add_run("\n")
            run_br.font.size = Pt(12)

    add_blank(doc)


def build_test_plan(doc, meta):
    """Section: 测试计划"""
    add_section_header(doc, "测试计划")

    start = meta.get("test_start_date", "")
    end = meta.get("test_end_date", "")
    testers = meta.get("testers", "")
    smoke = meta.get("smoke_test_date", "")
    sys_start = meta.get("system_test_start", "")
    sys_end = meta.get("system_test_end", "")

    if start or end:
        add_content(doc, f"测试时间：{start}-{end}")
    if testers:
        add_content(doc, f"测试人员：{testers}")
    if smoke:
        add_content(doc, f"冒烟测试时间：{smoke}")
    if sys_start or sys_end:
        add_content(doc, f"系统测试时间：{sys_start}-{sys_end}")

    add_blank(doc)


def build_test_environment(doc, meta):
    """Section: 测试环境及配置 (tab-separated key-values)."""
    add_section_header(doc, "测试环境及配置：")

    os_name = meta.get("os", "")
    device = meta.get("device_name", "")
    processor = meta.get("processor", "")
    ram = meta.get("ram", "")

    # Build environment info paragraph with tab-separated items
    lines = []
    if os_name:
        lines.append(f"系统：{os_name}")
    if device:
        lines.append(f"设备名称\t{device}")
    if processor:
        lines.append(f"处理器\t{processor}")
    if ram:
        lines.append(f"机带 RAM\t{ram}")

    if lines:
        text = "\n".join(lines)
        add_content(doc, text)
    else:
        add_content(doc, "（未配置）")

    add_blank(doc)


def build_test_basis(doc, meta):
    """Section: 测试依据"""
    add_section_header(doc, "测试依据：")
    basis = meta.get("test_basis", "根据开发提供的测试范围")
    add_content(doc, basis)
    add_blank(doc)


def build_test_cases(doc, data):
    """Section: 测试用例 (execution stats)."""
    add_section_header(doc, "测试用例：")

    tc = data.get("test_cases", {})
    total = tc.get("total", 0)
    executed = tc.get("executed", 0)
    exec_rate = tc.get("execution_rate", f"{executed}/{total}")

    if total > 0:
        add_content(doc, f"测试用例总数：{total}个")
        add_content(doc, f"已执行：{executed}个")
        add_content(doc, f"测试用例执行率：{exec_rate}")

        # Detailed case-by-case stats if provided
        passed = tc.get("passed")
        failed = tc.get("failed")
        if passed is not None:
            add_content(doc, f"通过：{passed}个")
        if failed is not None:
            add_content(doc, f"失败：{failed}个")
    else:
        add_content(doc, "（待补充）")

    add_blank(doc)


def build_test_defects(doc, data):
    """Section: 测试缺陷 (bug link + fix rate)."""
    add_section_header(doc, "测试缺陷")

    meta = data.get("report_meta", {})
    bugs = data.get("bugs", {})

    bug_url = meta.get("bug_system_url", "")
    bug_total = bugs.get("total", 0)
    bug_fixed = bugs.get("fixed", 0)
    fix_rate = bugs.get("fix_rate", f"{(bug_fixed / bug_total * 100):.0f}%" if bug_total > 0 else "N/A")

    # Bug system link
    if bug_url:
        add_content(doc, f"禅道bug链接：")
        add_content(doc, bug_url)

    # Bug statistics
    add_content(doc, f"测试bug总计：{bug_total}个，修复BUG：{bug_fixed}个")
    add_content(doc, f"bug修复率：{fix_rate}")

    # Detailed bug list if provided
    bug_details = bugs.get("details", [])
    if bug_details:
        add_blank(doc)
        add_content(doc, "BUG详细列表：", bold=True)
        for i, bug in enumerate(bug_details, 1):
            sev = bug.get("severity", "")
            status = bug.get("status", "")
            title = bug.get("title", "")
            bug_id = bug.get("id", "")
            add_content(doc, f"{i}. [{sev}][{status}] {bug_id} - {title}")

    add_blank(doc)


def build_risk_assessment(doc, meta):
    """Section: 风险评估"""
    add_section_header(doc, "风险评估")
    risk = meta.get("risk_assessment", "")
    if risk:
        add_content(doc, risk)
    else:
        add_content(doc, "（待评估）")
    add_blank(doc)


def build_doc_verification(doc, meta):
    """Section: 产品功能说明文档验证结果"""
    add_section_header(doc, "产品功能说明文档验证结果")

    doc_status = meta.get("doc_status", "未提供")
    clarity_status = meta.get("doc_clarity_status", "未提供")

    add_content(doc, f"产品功能说明文档链接或者文档\t{doc_status}")
    add_content(doc, f"根据产品功能说明文档操作是否清晰明了？{clarity_status}")
    add_blank(doc)


def build_test_results(doc, data):
    """Section: 测试结果 (execution summary + bug stats + conclusion)."""
    add_section_header(doc, "测试结果")

    tc = data.get("test_cases", {})
    bugs = data.get("bugs", {})
    meta = data.get("report_meta", {})

    total = tc.get("total", 0)
    executed = tc.get("executed", 0)
    exec_rate = tc.get("execution_rate", "100%")
    bug_total = bugs.get("total", 0)
    bug_fixed = bugs.get("fixed", 0)
    bug_fix_rate = bugs.get("fix_rate", "100%")
    summary = data.get("test_summary", meta.get("risk_assessment", ""))

    add_content(doc, f"测试用例执行情况：{total}个测试用例全部执行，")
    add_content(doc, f"测试用例执行率：{exec_rate}")
    add_content(doc, f"测试bug总计：{bug_total}个，修复BUG：{bug_fixed}个，BUG修复率：{bug_fix_rate}")
    if summary:
        add_content(doc, summary)

    add_blank(doc)


def build_conclusion(doc, data):
    """Section: 测试结论"""
    conclusion = data.get("conclusion", "")
    add_content(doc, f"测试结论：{conclusion}")


# ── document setup ───────────────────────────────────────────────────

def setup_document():
    """Create and configure the document matching the template."""
    doc = Document()

    # Page: A4
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.175)
        section.right_margin = Cm(3.175)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(14)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # Ensure List Paragraph style exists and has proper font
    if "List Paragraph" not in [s.name for s in doc.styles]:
        lp_style = doc.styles.add_style("List Paragraph", 1)
    lp_style = doc.styles["List Paragraph"]
    lp_style.font.name = "宋体"
    lp_style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    return doc


# ── main generation ──────────────────────────────────────────────────

def generate_report(input_data, output_path):
    """Generate the complete test report as DOCX."""
    doc = setup_document()
    meta = input_data.get("report_meta", {})

    # 1. Title
    build_title(doc, meta)
    add_blank(doc)

    # 2. 需求地址
    build_requirement_url(doc, meta)

    # 3. 测试范围
    build_test_scope(doc, input_data)

    # 4. 测试计划
    build_test_plan(doc, meta)

    # 5. 测试环境及配置
    build_test_environment(doc, meta)

    # 6. 测试依据
    build_test_basis(doc, meta)

    # 7. 测试用例
    build_test_cases(doc, input_data)

    # 8. 测试缺陷
    build_test_defects(doc, input_data)

    # 9. 风险评估
    build_risk_assessment(doc, meta)

    # 10. 产品功能说明文档验证结果
    build_doc_verification(doc, meta)

    # 11. 测试结果
    build_test_results(doc, input_data)

    # 12. 测试结论
    build_conclusion(doc, input_data)

    # Save
    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(
        description="Generate a Word (.docx) test report from JSON data."
    )
    parser.add_argument("--input", help="Path to JSON input file")
    parser.add_argument("--output", required=True, help="Path to output .docx file")
    parser.add_argument("--json-stdin", action="store_true",
                        help="Read JSON from stdin instead of file")

    args = parser.parse_args()

    # Read input
    if args.json_stdin:
        raw = sys.stdin.read()
        input_data = json.loads(raw)
    elif args.input:
        with open(args.input, "r", encoding="utf-8") as f:
            input_data = json.load(f)
    else:
        print("[ERROR] Either --input or --json-stdin is required.")
        sys.exit(1)

    # Validate
    if "report_meta" not in input_data:
        print("[ERROR] Input JSON must contain 'report_meta' field.")
        sys.exit(1)

    # Generate
    output = generate_report(input_data, args.output)
    print(f"[OK] Report generated: {output}")


if __name__ == "__main__":
    main()
